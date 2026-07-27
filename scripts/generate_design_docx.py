"""
Karibu Credit — Technical Design Document Generator
====================================================
Generates 5 Workflow Diagrams + 7 Dashboard Wireframe Mockups using Matplotlib
and compiles them into a styled Word (.docx) document.

Usage:
    python scripts/generate_design_docx.py

Output:
    docs/Karibu_Credit_Technical_Design_v1.docx
"""

import os
import sys
import io
import numpy as np

import matplotlib
matplotlib.use("Agg")  # Non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ─── Global Styles ───────────────────────────────────────────────────────────
WHITE = "#FFFFFF"
BG_DARK = "#0A0D18"
SURFACE = "#0F111A"
BRAND = "#4F46E5"
ACCENT = "#0EA5E9"
SUCCESS = "#10B981"
WARNING = "#F59E0B"
DANGER = "#EF4444"
INFO = "#8B5CF6"
TEXT_PRIMARY = "#F8FAFC"
TEXT_MUTED = "#64748B"
BORDER = "#1E293B"

# Print-friendly palette (light background for Word readability)
P_BG = "#FFFFFF"
P_SURFACE = "#F8FAFC"
P_TEXT = "#0F172A"
P_TEXT_SEC = "#475569"
P_BRAND = "#4338CA"
P_ACCENT = "#0284C7"
P_SUCCESS = "#059669"
P_WARNING = "#D97706"
P_DANGER = "#DC2626"
P_INFO = "#7C3AED"
P_BORDER = "#E2E8F0"
P_CARD = "#F1F5F9"

DPI = 180

# ─── Helpers ─────────────────────────────────────────────────────────────────
def fig_to_bytes(fig) -> bytes:
    """Render a Matplotlib figure to PNG bytes."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=DPI, bbox_inches="tight",
                facecolor=fig.get_facecolor(), edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def draw_box(ax, x, y, w, h, text, color=P_BRAND, text_color=WHITE,
             fontsize=8, radius=0.3, alpha=1.0, bold=True):
    """Draw a rounded rectangle box with centered text."""
    box = FancyBboxPatch(
        (x - w/2, y - h/2), w, h,
        boxstyle=f"round,pad=0,rounding_size={radius}",
        facecolor=color, edgecolor="none", alpha=alpha,
        transform=ax.transData, zorder=3
    )
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight="bold" if bold else "normal",
            wrap=True, zorder=4)


def draw_diamond(ax, x, y, size, text, color=P_WARNING, text_color=P_TEXT):
    """Draw a decision diamond."""
    d = size / 2
    diamond = plt.Polygon(
        [(x, y+d), (x+d, y), (x, y-d), (x-d, y)],
        facecolor=color, edgecolor="#B45309", linewidth=1.2, zorder=3
    )
    ax.add_patch(diamond)
    ax.text(x, y, text, ha="center", va="center", fontsize=6.5,
            color=text_color, fontweight="bold", zorder=4)


def draw_arrow(ax, x1, y1, x2, y2, label="", color=P_TEXT_SEC):
    """Draw an arrow between two points with an optional label."""
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=1.4,
                        connectionstyle="arc3,rad=0"),
        zorder=2
    )
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx + 0.15, my + 0.15, label, fontsize=5.5,
                color=P_TEXT_SEC, fontstyle="italic", zorder=5)


def draw_circle(ax, x, y, r, text, color=P_BRAND, text_color=WHITE, fontsize=7):
    """Draw a circle node (start/end)."""
    circle = plt.Circle((x, y), r, facecolor=color, edgecolor="none", zorder=3)
    ax.add_patch(circle)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight="bold", zorder=4)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORKFLOW DIAGRAM 1: Loan Application & Disbursement Flow
# ═══════════════════════════════════════════════════════════════════════════════
def draw_workflow_1():
    fig, ax = plt.subplots(figsize=(11, 7))
    fig.set_facecolor(P_BG)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Workflow 1: Loan Application & Disbursement Flow",
                 fontsize=12, fontweight="bold", color=P_TEXT, pad=12)

    # Swimlane headers
    lanes = [("Customer", 1.2), ("Loan Officer", 3.3), ("Finance / System", 5.4),
             ("ZamuPay Gateway", 7.5)]
    for label, y_pos in lanes:
        ax.axhline(y=y_pos - 0.7, color=P_BORDER, linewidth=0.6, linestyle="--")
        ax.text(0.3, y_pos, label, fontsize=7, fontweight="bold",
                color=P_BRAND, rotation=0, va="center")

    # Nodes
    draw_circle(ax, 2.5, 1.2, 0.35, "Apply", P_BRAND)
    draw_box(ax, 4.5, 1.2, 2.2, 0.65, "POST /apply\n(Amount, Type)", P_ACCENT)
    draw_box(ax, 7.5, 1.2, 2.0, 0.65, "Save Loan\nstatus=PENDING", P_CARD, P_TEXT, alpha=0.9)

    draw_box(ax, 3.5, 3.3, 2.2, 0.65, "Review Queue\nGET /loans?pending", "#E0E7FF", P_BRAND)
    draw_box(ax, 6.5, 3.3, 2.0, 0.65, "POST /{id}/approve\nSet Schedule", P_SUCCESS)

    draw_diamond(ax, 9.5, 5.4, 1.1, "Lump\nSum?", P_WARNING)
    draw_box(ax, 6.0, 5.4, 2.2, 0.65, "Status=APPROVED\n(Hold Funds)", "#FEF3C7", P_TEXT)
    draw_box(ax, 12.0, 5.4, 2.0, 0.65, "Full Payout\nvia ZamuPay B2C", P_SUCCESS)

    draw_box(ax, 6.0, 7.5, 2.4, 0.65, "Tranche Request\nPOST /disburse_tranche", P_ACCENT)
    draw_box(ax, 9.5, 7.5, 2.0, 0.65, "Partial Payout\nvia ZamuPay B2C", "#DBEAFE", P_TEXT)
    draw_box(ax, 12.5, 7.5, 1.8, 0.65, "Record Tx\n+ Update Loan", P_CARD, P_TEXT)

    # Arrows
    draw_arrow(ax, 2.85, 1.2, 3.4, 1.2)
    draw_arrow(ax, 5.6, 1.2, 6.5, 1.2)
    draw_arrow(ax, 7.5, 1.55, 3.5, 2.95, "Officer picks up")
    draw_arrow(ax, 4.6, 3.3, 5.5, 3.3)
    draw_arrow(ax, 7.5, 3.3, 9.5, 4.85)

    draw_arrow(ax, 10.05, 5.4, 11.0, 5.4, "Yes")
    draw_arrow(ax, 9.5, 4.85, 6.0, 4.85)
    ax.text(8.5, 4.6, "No (Stage-wise)", fontsize=5.5, color=P_WARNING, fontstyle="italic")

    draw_arrow(ax, 6.0, 5.05, 6.0, 7.15, "Later...")
    draw_arrow(ax, 7.2, 7.5, 8.5, 7.5)
    draw_arrow(ax, 10.5, 7.5, 11.6, 7.5)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORKFLOW DIAGRAM 2: M-Pesa Repayment & Auto-Clearing
# ═══════════════════════════════════════════════════════════════════════════════
def draw_workflow_2():
    fig, ax = plt.subplots(figsize=(11, 6))
    fig.set_facecolor(P_BG)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Workflow 2: M-Pesa C2B Repayment & Auto-Clearing Flow",
                 fontsize=12, fontweight="bold", color=P_TEXT, pad=12)

    draw_circle(ax, 1.5, 6.5, 0.4, "Customer\nPays", P_SUCCESS)
    draw_box(ax, 4.0, 6.5, 2.5, 0.7, "M-Pesa C2B Paybill\n(Account = Loan ID)", "#DCFCE7", P_TEXT)
    draw_box(ax, 7.5, 6.5, 2.5, 0.7, "POST /c2b/validation\n→ {ResultCode: 0}", P_ACCENT)

    draw_box(ax, 4.0, 4.5, 2.5, 0.7, "POST /c2b/confirmation\n(JSON Payload)", P_BRAND)
    draw_box(ax, 7.5, 4.5, 2.2, 0.7, "Find Loan by\nAccount Ref", P_CARD, P_TEXT)
    draw_box(ax, 10.8, 4.5, 2.2, 0.7, "allocate_repayment()\nPenalties → Interest\n→ Principal", P_INFO, WHITE)

    draw_diamond(ax, 7.5, 2.5, 1.0, "Balance\n≤ 0?", P_WARNING)
    draw_box(ax, 4.0, 2.5, 2.0, 0.65, "Record Payment\n+ Transaction", P_SUCCESS)
    draw_box(ax, 10.8, 2.5, 2.2, 0.65, "Status = CLEARED\nLoan closed", "#DCFCE7", P_TEXT)

    draw_box(ax, 7.5, 0.8, 2.2, 0.55, "Commit DB Session\nReturn {ResultCode: 0}", P_CARD, P_TEXT)

    # Arrows
    draw_arrow(ax, 1.9, 6.5, 2.75, 6.5)
    draw_arrow(ax, 5.25, 6.5, 6.25, 6.5)
    draw_arrow(ax, 7.5, 6.15, 4.0, 4.85, "Safaricom confirms")
    draw_arrow(ax, 5.25, 4.5, 6.4, 4.5)
    draw_arrow(ax, 8.6, 4.5, 9.7, 4.5)

    draw_arrow(ax, 10.8, 4.15, 7.5, 3.0)
    draw_arrow(ax, 7.0, 2.5, 5.0, 2.5, "Always")
    draw_arrow(ax, 8.0, 2.5, 9.7, 2.5, "Yes")
    draw_arrow(ax, 7.5, 2.0, 7.5, 1.1)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORKFLOW DIAGRAM 3: Nightly Penalties & Defaults (Cron Job)
# ═══════════════════════════════════════════════════════════════════════════════
def draw_workflow_3():
    fig, ax = plt.subplots(figsize=(10, 7))
    fig.set_facecolor(P_BG)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Workflow 3: Nightly Penalties & Defaults Cron Job",
                 fontsize=12, fontweight="bold", color=P_TEXT, pad=12)

    draw_circle(ax, 6, 8.2, 0.45, "Midnight\nCron", P_BRAND)
    draw_box(ax, 6, 7.0, 3.0, 0.65, "Query DB: Active Loans\npast due_date", P_ACCENT)
    draw_diamond(ax, 6, 5.5, 1.1, "Overdue\nloans?", P_WARNING)

    draw_box(ax, 2.5, 5.5, 1.8, 0.55, "Sleep until\ntomorrow", P_CARD, P_TEXT)

    draw_box(ax, 6, 4.0, 2.5, 0.6, "Loop: each overdue loan", "#E0E7FF", P_BRAND)
    draw_box(ax, 6, 2.8, 3.0, 0.65, "Calculate Outstanding\n(payable − paid + penalties)", P_INFO, WHITE)
    draw_diamond(ax, 6, 1.5, 0.9, "Owed\n> 0?", P_WARNING)

    draw_box(ax, 9.5, 1.5, 2.2, 0.55, "Apply 10% Penalty\nSet DEFAULTED", P_DANGER, WHITE)
    draw_box(ax, 9.5, 0.4, 2.0, 0.5, "Save Transaction\nto DB", P_SUCCESS)

    draw_box(ax, 3.0, 1.5, 1.5, 0.5, "Skip\n(No debt)", P_CARD, P_TEXT)

    # Arrows
    draw_arrow(ax, 6, 7.75, 6, 7.3)
    draw_arrow(ax, 6, 6.65, 6, 6.05)
    draw_arrow(ax, 5.45, 5.5, 3.4, 5.5, "No")
    draw_arrow(ax, 6, 5.0, 6, 4.3, "Yes")
    draw_arrow(ax, 6, 3.7, 6, 3.1)
    draw_arrow(ax, 6, 2.45, 6, 1.95)
    draw_arrow(ax, 6.45, 1.5, 8.4, 1.5, "Yes")
    draw_arrow(ax, 5.55, 1.5, 3.75, 1.5, "No")
    draw_arrow(ax, 9.5, 1.22, 9.5, 0.65)

    # Loop-back arrow (hand-drawn style)
    ax.annotate("", xy=(6, 3.7), xytext=(11.0, 0.4),
                arrowprops=dict(arrowstyle="-|>", color=P_TEXT_SEC, lw=1.0,
                                connectionstyle="arc3,rad=-0.4"), zorder=2)
    ax.text(11.3, 2.0, "Next\nLoan", fontsize=5.5, color=P_TEXT_SEC, fontstyle="italic")

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORKFLOW DIAGRAM 4: Loan Top-Up & Settlement Flow
# ═══════════════════════════════════════════════════════════════════════════════
def draw_workflow_4():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    fig.set_facecolor(P_BG)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Workflow 4: Loan Top-Up & Settlement Flow",
                 fontsize=12, fontweight="bold", color=P_TEXT, pad=12)

    draw_circle(ax, 1.5, 6.5, 0.4, "Borrower\nRequest", P_BRAND)
    draw_box(ax, 4.0, 6.5, 2.5, 0.7, "POST /loans/{id}/top-up\n(additional_amount)", P_ACCENT)
    draw_diamond(ax, 7.5, 6.5, 1.1, "Active\nLoan?", P_WARNING)

    draw_box(ax, 10.5, 6.5, 1.8, 0.55, "Reject: No\nactive loan", P_DANGER, WHITE)
    draw_box(ax, 7.5, 4.5, 2.8, 0.7, "Calculate:\nold_balance + top_up_amount\n= new_principal", P_INFO, WHITE)
    draw_box(ax, 3.5, 4.5, 2.5, 0.7, "Check repayment\nhistory ≥ 50% paid", "#FEF3C7", P_TEXT)

    draw_diamond(ax, 3.5, 2.8, 1.0, "Eligible?", P_WARNING)
    draw_box(ax, 7.5, 2.8, 2.8, 0.65, "Recalculate schedule\nMerge balances, new term", P_SUCCESS)
    draw_box(ax, 11.5, 2.8, 2.2, 0.65, "Disburse difference\nvia ZamuPay B2C", P_ACCENT)
    draw_box(ax, 1.0, 2.8, 1.5, 0.5, "Reject\ntop-up", P_DANGER, WHITE)

    draw_box(ax, 7.5, 1.0, 3.0, 0.55, "Audit Trail + SMS Notification\nstatus=ACTIVE (merged)", P_CARD, P_TEXT)

    # Arrows
    draw_arrow(ax, 1.9, 6.5, 2.75, 6.5)
    draw_arrow(ax, 5.25, 6.5, 6.95, 6.5)
    draw_arrow(ax, 8.05, 6.5, 9.6, 6.5, "No")
    draw_arrow(ax, 7.5, 5.95, 7.5, 4.85, "Yes")
    draw_arrow(ax, 6.1, 4.5, 4.75, 4.5)
    draw_arrow(ax, 3.5, 4.15, 3.5, 3.3)
    draw_arrow(ax, 4.0, 2.8, 6.1, 2.8, "Yes")
    draw_arrow(ax, 3.0, 2.8, 1.75, 2.8, "No")
    draw_arrow(ax, 8.9, 2.8, 10.4, 2.8)
    draw_arrow(ax, 7.5, 2.45, 7.5, 1.3)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORKFLOW DIAGRAM 5: Group Lending Joint Liability Flow
# ═══════════════════════════════════════════════════════════════════════════════
def draw_workflow_5():
    fig, ax = plt.subplots(figsize=(11, 7))
    fig.set_facecolor(P_BG)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Workflow 5: Group Lending — Joint Liability Disbursement",
                 fontsize=12, fontweight="bold", color=P_TEXT, pad=12)

    draw_circle(ax, 1.5, 7.5, 0.4, "Group\nChair", P_BRAND)
    draw_box(ax, 4.0, 7.5, 2.5, 0.7, "POST /groups/create\n(name, members)", P_ACCENT)
    draw_box(ax, 7.5, 7.5, 2.2, 0.65, "Create Group\nAssign Members", P_SUCCESS)

    draw_box(ax, 4.0, 5.8, 2.5, 0.7, "POST /groups/apply\n(group_id, amount)", P_BRAND)
    draw_diamond(ax, 7.5, 5.8, 1.1, "All KYC\nVerified?", P_WARNING)
    draw_box(ax, 10.5, 5.8, 2.0, 0.55, "Reject: KYC\nincomplete", P_DANGER, WHITE)

    draw_box(ax, 7.5, 4.0, 2.8, 0.7, "Split principal across\nmembers (equal share)", P_INFO, WHITE)
    draw_box(ax, 3.5, 4.0, 2.5, 0.7, "Each member gets\nindividual schedule", "#E0E7FF", P_BRAND)
    draw_box(ax, 7.5, 2.2, 2.5, 0.65, "Disburse to each\nmember via B2C", P_ACCENT)

    draw_box(ax, 3.5, 2.2, 2.5, 0.65, "Joint Liability:\nIf 1 defaults, group\ncovers shortfall", "#FEF3C7", P_TEXT)
    draw_box(ax, 11.0, 2.2, 2.2, 0.65, "Track group\nrepayment rate", P_SUCCESS)

    draw_box(ax, 7.5, 0.5, 3.0, 0.55, "Audit Trail: Group loan lifecycle logged\nSMS to all members", P_CARD, P_TEXT)

    # Arrows
    draw_arrow(ax, 1.9, 7.5, 2.75, 7.5)
    draw_arrow(ax, 5.25, 7.5, 6.4, 7.5)
    draw_arrow(ax, 7.5, 7.15, 4.0, 6.15, "Chair applies")
    draw_arrow(ax, 5.25, 5.8, 6.95, 5.8)
    draw_arrow(ax, 8.05, 5.8, 9.5, 5.8, "No")
    draw_arrow(ax, 7.5, 5.25, 7.5, 4.35, "Yes")
    draw_arrow(ax, 6.1, 4.0, 4.75, 4.0)
    draw_arrow(ax, 7.5, 3.65, 7.5, 2.55)
    draw_arrow(ax, 6.25, 2.2, 4.75, 2.2)
    draw_arrow(ax, 8.75, 2.2, 9.9, 2.2)
    draw_arrow(ax, 7.5, 1.85, 7.5, 0.8)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD MOCKUP HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def draw_kpi_card(ax, x, y, w, h, label, value, sub, accent_color):
    """Draw a KPI metric card on the axes."""
    box = FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0,rounding_size=0.08",
        facecolor=P_SURFACE, edgecolor=P_BORDER, linewidth=0.8, zorder=3
    )
    ax.add_patch(box)
    # Accent top bar
    ax.plot([x+0.02, x+w-0.02], [y+h-0.01, y+h-0.01],
            color=accent_color, linewidth=3, solid_capstyle="round", zorder=4)
    ax.text(x + w/2, y + h*0.7, label, ha="center", va="center",
            fontsize=5, color=TEXT_MUTED, fontweight="bold", zorder=5)
    ax.text(x + w/2, y + h*0.42, value, ha="center", va="center",
            fontsize=9, color=TEXT_PRIMARY, fontweight="bold", zorder=5)
    ax.text(x + w/2, y + h*0.15, sub, ha="center", va="center",
            fontsize=4.5, color=TEXT_MUTED, zorder=5)


def draw_mini_bar(ax_inset, data, colors, labels=None):
    """Draw a small bar chart in an inset axes."""
    ax_inset.set_facecolor(SURFACE)
    bars = ax_inset.bar(range(len(data)), data, color=colors, width=0.6, zorder=3)
    ax_inset.set_xticks(range(len(data)))
    if labels:
        ax_inset.set_xticklabels(labels, fontsize=4, color=TEXT_MUTED, rotation=30)
    else:
        ax_inset.set_xticklabels([])
    ax_inset.tick_params(axis="y", labelsize=4, colors=TEXT_MUTED)
    ax_inset.spines["top"].set_visible(False)
    ax_inset.spines["right"].set_visible(False)
    ax_inset.spines["left"].set_color(BORDER)
    ax_inset.spines["bottom"].set_color(BORDER)
    for bar in bars:
        bar.set_edgecolor("none")


def draw_mini_line(ax_inset, data, color=ACCENT, fill=True):
    """Draw a small line chart in an inset axes."""
    ax_inset.set_facecolor(SURFACE)
    x = range(len(data))
    ax_inset.plot(x, data, color=color, linewidth=1.5, zorder=3)
    if fill:
        ax_inset.fill_between(x, data, alpha=0.15, color=color, zorder=2)
    ax_inset.set_xticks([])
    ax_inset.tick_params(axis="y", labelsize=4, colors=TEXT_MUTED)
    ax_inset.spines["top"].set_visible(False)
    ax_inset.spines["right"].set_visible(False)
    ax_inset.spines["left"].set_color(BORDER)
    ax_inset.spines["bottom"].set_color(BORDER)


def draw_mini_pie(ax_inset, data, colors, labels=None):
    """Draw a small pie chart in an inset axes."""
    ax_inset.set_facecolor(SURFACE)
    wedges, texts = ax_inset.pie(data, colors=colors, startangle=90,
                                  wedgeprops=dict(width=0.4, edgecolor=SURFACE))
    if labels:
        ax_inset.legend(wedges, labels, loc="center", fontsize=3.5,
                        frameon=False, labelcolor=TEXT_MUTED)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 1: CEO / Managing Director
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_1():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "CEO / Managing Director Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)
    ax.text(6, 7.35, "Executive overview • Portfolio health • Revenue metrics",
            ha="center", fontsize=7, color=TEXT_MUTED)

    # KPI row
    kpis = [
        ("Total Portfolio", "KES 64.2M", "+12.4% MoM", BRAND),
        ("Monthly Revenue", "KES 22.4M", "+8.6% target", SUCCESS),
        ("Net Profit Margin", "24.8%", "Within target", ACCENT),
        ("Active Customers", "358", "+18 this week", INFO),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Charts row
    # Bar chart: Disbursements vs Collections
    ax_bar = ax.inset_axes([0.04, 0.15, 0.44, 0.42])
    months = ["Dec", "Jan", "Feb", "Mar", "Apr", "May"]
    disb = [12.4, 15.1, 18.2, 14.8, 21.3, 26.0]
    coll = [8.5, 11.2, 13.6, 12.1, 16.5, 22.4]
    x = np.arange(len(months))
    ax_bar.set_facecolor(SURFACE)
    ax_bar.bar(x - 0.2, disb, 0.35, color=BRAND, label="Disbursements", zorder=3)
    ax_bar.bar(x + 0.2, coll, 0.35, color=SUCCESS, label="Collections", zorder=3)
    ax_bar.set_xticks(x)
    ax_bar.set_xticklabels(months, fontsize=5, color=TEXT_MUTED)
    ax_bar.tick_params(axis="y", labelsize=5, colors=TEXT_MUTED)
    ax_bar.legend(fontsize=5, frameon=False, labelcolor=TEXT_MUTED)
    ax_bar.set_title("Disbursements vs Collections (KES M)", fontsize=6,
                     color=TEXT_PRIMARY, fontweight="bold")
    ax_bar.spines["top"].set_visible(False)
    ax_bar.spines["right"].set_visible(False)
    ax_bar.spines["left"].set_color(BORDER)
    ax_bar.spines["bottom"].set_color(BORDER)

    # Doughnut: Portfolio mix
    ax_pie = ax.inset_axes([0.55, 0.15, 0.4, 0.42])
    ax_pie.set_facecolor(SURFACE)
    mix_data = [45, 30, 15, 10]
    mix_colors = [BRAND, SUCCESS, ACCENT, WARNING]
    mix_labels = ["Logbook", "SME", "Agri", "Personal"]
    wedges, _ = ax_pie.pie(mix_data, colors=mix_colors, startangle=90,
                           wedgeprops=dict(width=0.45, edgecolor=SURFACE))
    ax_pie.legend(wedges, mix_labels, loc="center", fontsize=5,
                  frameon=False, labelcolor=TEXT_MUTED)
    ax_pie.set_title("Portfolio by Product Mix", fontsize=6,
                     color=TEXT_PRIMARY, fontweight="bold")

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 2: CFO / Finance
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_2():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "CFO / Finance Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)
    ax.text(6, 7.35, "Revenue streams • Daraja settlement audit • Provision reserves",
            ha="center", fontsize=7, color=TEXT_MUTED)

    kpis = [
        ("Interest Income", "KES 14.8M", "MTD", SUCCESS),
        ("Processing Fees", "KES 3.25M", "MTD", ACCENT),
        ("Gross Profit", "KES 18.9M", "MTD", BRAND),
        ("Provision Reserve", "KES 8.45M", "Defaulted value", DANGER),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Settlement audit table mockup
    table_y = 4.8
    ax.text(1.0, table_y, "Daraja B2C / C2B Settlement Audit", fontsize=8,
            fontweight="bold", color=TEXT_PRIMARY)
    headers = ["Transaction Pool", "Tx Count", "Total Value", "SLA Status"]
    rows = [
        ("B2C Disbursed Sent", "182", "KES 9,840,000", "MATCHED"),
        ("B2C Confirmed Callback", "180", "KES 9,740,000", "MATCHED"),
        ("B2C Pending Callback", "2", "KES 100,000", "PENDING"),
        ("C2B Paybill MTD", "1,420", "KES 22,400,000", "MATCHED"),
    ]
    for j, h in enumerate(headers):
        ax.text(1.0 + j*2.7, table_y - 0.4, h, fontsize=5.5,
                fontweight="bold", color=TEXT_MUTED)
    for i, row in enumerate(rows):
        y = table_y - 0.8 - i*0.35
        status_color = SUCCESS if row[3] == "MATCHED" else WARNING
        for j, cell in enumerate(row):
            c = status_color if j == 3 else TEXT_PRIMARY
            fw = "bold" if j == 3 else "normal"
            ax.text(1.0 + j*2.7, y, cell, fontsize=5.5, color=c, fontweight=fw)
        ax.axhline(y=y-0.12, xmin=0.06, xmax=0.94, color=BORDER, linewidth=0.3)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 3: Branch Manager
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_3():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "Branch Manager Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)

    kpis = [
        ("Branch Active Loans", "142", "As of today", BRAND),
        ("Branch Book Value", "KES 24.5M", "+5.2% MoM", ACCENT),
        ("Branch PAR 30", "3.82%", "Within target", SUCCESS),
        ("Escalations", "3 Pending", "Awaiting review", WARNING),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Officer performance table
    ax.text(1.0, 4.8, "Officer Performance Board", fontsize=8,
            fontweight="bold", color=TEXT_PRIMARY)
    headers = ["Loan Officer", "Apps", "Avg TAT", "PAR 30"]
    officers = [
        ("David Kipkorir", "48", "14.5 hrs", "2.4%"),
        ("Sarah Wambui", "42", "16.2 hrs", "3.1%"),
        ("Mark Omwansa", "35", "21.0 hrs", "6.8%"),
        ("Grace Mutheu", "28", "25.5 hrs", "9.2%"),
    ]
    for j, h in enumerate(headers):
        ax.text(1.0 + j*2.5, 4.4, h, fontsize=5.5,
                fontweight="bold", color=TEXT_MUTED)
    for i, row in enumerate(officers):
        y = 4.0 - i*0.35
        for j, cell in enumerate(row):
            c = DANGER if j == 3 and float(cell.replace("%","")) > 8 else TEXT_PRIMARY
            ax.text(1.0 + j*2.5, y, cell, fontsize=5.5, color=c)
        ax.axhline(y=y-0.12, xmin=0.06, xmax=0.85, color=BORDER, linewidth=0.3)

    # PAR breakdown bars
    ax_par = ax.inset_axes([0.55, 0.08, 0.4, 0.35])
    par_labels = ["1–7d", "8–30d", "31–90d", "90+d"]
    par_data = [22, 8, 3, 1]
    par_colors = [SUCCESS, WARNING, DANGER, "#991B1B"]
    draw_mini_bar(ax_par, par_data, par_colors, par_labels)
    ax_par.set_title("PAR Aging Breakdown (Accounts)", fontsize=5,
                     color=TEXT_PRIMARY, fontweight="bold")

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 4: Loan Officer
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_4():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "Loan Officer Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)
    ax.text(6, 7.35, "Application queue • Appraisal checklist • Approval workflow",
            ha="center", fontsize=7, color=TEXT_MUTED)

    kpis = [
        ("My Queue", "5 apps", "Pending review", BRAND),
        ("Approval Rate", "78.2%", "Above target", SUCCESS),
        ("Avg Review TAT", "16.4 hrs", "SLA: 24h", WARNING),
        ("Collection Rate", "96.8%", "Origination", ACCENT),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Application queue
    ax.text(1.0, 4.8, "Application Queue (Click to Appraise)", fontsize=8,
            fontweight="bold", color=TEXT_PRIMARY)
    headers = ["Reference", "Borrower", "Product", "Amount"]
    queue = [
        ("LAF-2026-0012", "Grace Mwangi", "Logbook", "KES 500,000"),
        ("LAF-2026-0013", "Hassan Ibrahim", "SME", "KES 250,000"),
        ("LAF-2026-0014", "David Omondi", "Salary", "KES 80,000"),
    ]
    for j, h in enumerate(headers):
        ax.text(1.0 + j*2.5, 4.4, h, fontsize=5.5,
                fontweight="bold", color=TEXT_MUTED)
    for i, row in enumerate(queue):
        y = 4.0 - i*0.35
        for j, cell in enumerate(row):
            c = BRAND if j == 0 else TEXT_PRIMARY
            ax.text(1.0 + j*2.5, y, cell, fontsize=5.5, color=c,
                    fontweight="bold" if j == 0 else "normal")

    # Checklist panel
    ax.text(7.5, 2.8, "⬜ Appraisal Checklist", fontsize=7,
            fontweight="bold", color=TEXT_PRIMARY)
    checks = ["IPRS National ID Verified", "KRA Tax PIN validated",
              "TransUnion CRB clearance", "Payslip & Bank Statements",
              "Borrower interview note"]
    for i, item in enumerate(checks):
        ax.text(7.8, 2.4 - i*0.3, f"☐  {item}", fontsize=5.5, color=TEXT_MUTED)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 5: Collections Officer
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_5():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "Collections Officer Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)

    kpis = [
        ("Arrears Cases", "14 active", "Follow-up required", DANGER),
        ("Recovery Target", "KES 8.45M", "78% collected", ACCENT),
        ("Restructured", "12 borrowers", "This month", BRAND),
        ("CRB Submissions", "2 negative", "Listings filed", WARNING),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Kanban-style columns
    columns = [
        ("1–30 Days", SUCCESS, [("Grace Mwangi", "KES 15,000"), ("John Kamau", "KES 4,500")]),
        ("31–60 Days", WARNING, [("Hassan Ibrahim", "KES 35,000")]),
        ("60+ Days (Legal)", DANGER, [("Mercy Achieng", "KES 110,000")]),
    ]
    for col_i, (title, color, cards) in enumerate(columns):
        x = 0.8 + col_i * 3.8
        # Column header
        box = FancyBboxPatch(
            (x, 4.5), 3.4, 0.45, boxstyle="round,pad=0,rounding_size=0.06",
            facecolor=color, edgecolor="none", alpha=0.2, zorder=3
        )
        ax.add_patch(box)
        ax.text(x + 1.7, 4.72, title, ha="center", fontsize=7,
                fontweight="bold", color=color, zorder=4)
        ax.text(x + 3.1, 4.72, str(len(cards)), ha="center", fontsize=6,
                fontweight="bold", color=color, zorder=4)

        for card_i, (name, amount) in enumerate(cards):
            cy = 4.0 - card_i * 0.7
            card = FancyBboxPatch(
                (x + 0.1, cy - 0.2), 3.2, 0.55,
                boxstyle="round,pad=0,rounding_size=0.05",
                facecolor=SURFACE, edgecolor=BORDER, linewidth=0.5, zorder=3
            )
            ax.add_patch(card)
            ax.text(x + 0.3, cy + 0.1, name, fontsize=5.5,
                    fontweight="bold", color=TEXT_PRIMARY, zorder=4)
            ax.text(x + 0.3, cy - 0.1, f"Outstanding: {amount}", fontsize=5,
                    color=TEXT_MUTED, zorder=4)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 6: Compliance Officer
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_6():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "Compliance Officer Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)
    ax.text(6, 7.35, "CBK regulatory tracker • AML flags • KYC verification queue",
            ha="center", fontsize=7, color=TEXT_MUTED)

    # CBK Regulatory filings
    ax.text(1.0, 6.5, "CBK Regulatory Report Tracker", fontsize=8,
            fontweight="bold", color=TEXT_PRIMARY)
    headers = ["CBK Filing", "Due Date", "Auto-Draft", "Status"]
    filings = [
        ("Monthly Loan Returns", "May 31", "Completed", "DRAFT"),
        ("NPL & Credit Loss", "May 31", "Completed", "DRAFT"),
        ("Consumer Complaints", "Jun 30", "In Progress", "PENDING"),
        ("Quarterly Prudential", "Jun 30", "Manual Seeded", "PENDING"),
        ("Annual AML Board", "Dec 31", "Partial", "PENDING"),
    ]
    for j, h in enumerate(headers):
        ax.text(1.0 + j*2.7, 6.1, h, fontsize=5.5,
                fontweight="bold", color=TEXT_MUTED)
    for i, row in enumerate(filings):
        y = 5.7 - i*0.35
        for j, cell in enumerate(row):
            c = SUCCESS if cell in ("Completed","DRAFT") else WARNING if cell in ("In Progress","PENDING") else TEXT_PRIMARY
            ax.text(1.0 + j*2.7, y, cell, fontsize=5.5, color=c)
        ax.axhline(y=y-0.12, xmin=0.06, xmax=0.94, color=BORDER, linewidth=0.3)

    # AML section
    ax.text(1.0, 3.5, "AML Transaction Flags (≥ KES 300,000)", fontsize=7,
            fontweight="bold", color=WARNING)
    ax.text(1.0, 3.1, "3 flagged transactions requiring compliance review",
            fontsize=5.5, color=TEXT_MUTED)

    # KYC Queue
    ax.text(6.5, 3.5, "KYC Verification Queue", fontsize=7,
            fontweight="bold", color=INFO)
    ax.text(6.5, 3.1, "48-hour SLA for onboarding verification",
            fontsize=5.5, color=TEXT_MUTED)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD 7: Credit Scoring Engine
# ═══════════════════════════════════════════════════════════════════════════════
def draw_dashboard_7():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.set_facecolor(BG_DARK)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor(BG_DARK)
    ax.text(6, 7.7, "Credit Scoring Engine Dashboard", ha="center",
            fontsize=14, fontweight="bold", color=TEXT_PRIMARY)

    kpis = [
        ("Avg Score", "672 / 1000", "GOOD", SUCCESS),
        ("KYC Rate", "92.4%", "Above target", BRAND),
        ("CRB Listing", "4.2%", "Monitor closely", DANGER),
        ("Highest Score", "910", "Top borrower", ACCENT),
    ]
    for i, (label, value, sub, color) in enumerate(kpis):
        draw_kpi_card(ax, 0.3 + i*2.9, 5.8, 2.6, 1.2, label, value, sub, color)

    # Score distribution histogram
    ax_hist = ax.inset_axes([0.04, 0.15, 0.44, 0.45])
    ax_hist.set_facecolor(SURFACE)
    bins_labels = ["200–300", "300–499", "500–649", "650–799", "800–1000"]
    counts = [12, 45, 92, 148, 61]
    hist_colors = [DANGER, WARNING, ACCENT, BRAND, SUCCESS]
    ax_hist.bar(range(len(counts)), counts, color=hist_colors, width=0.65, zorder=3)
    ax_hist.set_xticks(range(len(bins_labels)))
    ax_hist.set_xticklabels(bins_labels, fontsize=4.5, color=TEXT_MUTED, rotation=20)
    ax_hist.tick_params(axis="y", labelsize=4, colors=TEXT_MUTED)
    ax_hist.set_title("Score Distribution (Borrower Count)", fontsize=6,
                      color=TEXT_PRIMARY, fontweight="bold")
    ax_hist.spines["top"].set_visible(False)
    ax_hist.spines["right"].set_visible(False)
    ax_hist.spines["left"].set_color(BORDER)
    ax_hist.spines["bottom"].set_color(BORDER)

    # Eligibility matrix
    ax.text(7.0, 4.8, "Eligibility Matrix", fontsize=8,
            fontweight="bold", color=TEXT_PRIMARY)
    tiers = [
        ("Excellent", "800–1000", "All products", "KES 1M", SUCCESS),
        ("Good", "650–799", "Standard terms", "KES 500K", SUCCESS),
        ("Fair", "500–649", "Collateral req.", "KES 200K", WARNING),
        ("Poor", "300–499", "Micro-loans only", "KES 30K", DANGER),
        ("Very Poor", "<300", "Auto-rejected", "KES 0", DANGER),
    ]
    headers = ["Tier", "Range", "Eligibility", "Ceiling"]
    for j, h in enumerate(headers):
        ax.text(7.0 + j*1.4, 4.4, h, fontsize=5, fontweight="bold", color=TEXT_MUTED)
    for i, (tier, rng, elig, ceil, color) in enumerate(tiers):
        y = 4.0 - i*0.35
        ax.text(7.0, y, tier, fontsize=5.5, fontweight="bold", color=color)
        ax.text(8.4, y, rng, fontsize=5.5, color=TEXT_PRIMARY)
        ax.text(9.8, y, elig, fontsize=5, color=TEXT_MUTED)
        ax.text(11.2, y, ceil, fontsize=5.5, fontweight="bold", color=TEXT_PRIMARY)

    return fig_to_bytes(fig)


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT COMPILATION
# ═══════════════════════════════════════════════════════════════════════════════
def compile_document():
    """Compile all diagrams into a styled Word document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Title page
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KARIBU CREDIT")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Dashboard Visual Mockups\n& Workflow Diagram Implementation Guide")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("7 Role Dashboards  ·  5 Workflow Diagrams  ·  Full Implementation Notes")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Karibu Credit Ltd  ·  CBK NDTCP Regulated  ·  Version 1.0  ·  May 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ─── Technology Stack page ───
    doc.add_heading("How the Diagrams Were Generated", level=1)
    doc.add_paragraph(
        "Every workflow diagram and dashboard wireframe in this document was produced "
        "entirely in Python using Matplotlib — no design tool, no external assets, no SVG editor. "
        "Each diagram is rendered as a high-resolution PNG (180 DPI) and embedded directly into "
        "this Word document using the python-docx library."
    )

    tech_table = doc.add_table(rows=6, cols=3)
    tech_table.style = "Light List Accent 1"
    headers = ["Library", "Version", "Role"]
    for i, h in enumerate(headers):
        tech_table.rows[0].cells[i].text = h
    tech_data = [
        ("matplotlib", "3.10.x", "Main drawing engine — axes, patches, text, arrows, charts"),
        ("matplotlib.patches", "Built-in", "FancyBboxPatch for rounded boxes, FancyArrowPatch for arrows"),
        ("numpy", "2.4.x", "Data arrays for chart series, score distributions, trend lines"),
        ("Pillow (PIL)", "12.x", "PNG encoding — required by matplotlib Agg backend"),
        ("python-docx", "1.2.x", "Embedding PNG bytes into Word .docx via InlineShape"),
    ]
    for i, (lib, ver, role) in enumerate(tech_data):
        row = tech_table.rows[i + 1]
        row.cells[0].text = lib
        row.cells[1].text = ver
        row.cells[2].text = role

    doc.add_page_break()

    # ─── Generate and embed all diagrams ───
    print("Generating Workflow Diagram 1: Loan Application & Disbursement...")
    wf1 = draw_workflow_1()
    print("Generating Workflow Diagram 2: M-Pesa Repayment & Auto-Clearing...")
    wf2 = draw_workflow_2()
    print("Generating Workflow Diagram 3: Nightly Penalties & Defaults...")
    wf3 = draw_workflow_3()
    print("Generating Workflow Diagram 4: Loan Top-Up & Settlement...")
    wf4 = draw_workflow_4()
    print("Generating Workflow Diagram 5: Group Lending Joint Liability...")
    wf5 = draw_workflow_5()

    print("Generating Dashboard 1: CEO / Managing Director...")
    db1 = draw_dashboard_1()
    print("Generating Dashboard 2: CFO / Finance...")
    db2 = draw_dashboard_2()
    print("Generating Dashboard 3: Branch Manager...")
    db3 = draw_dashboard_3()
    print("Generating Dashboard 4: Loan Officer...")
    db4 = draw_dashboard_4()
    print("Generating Dashboard 5: Collections Officer...")
    db5 = draw_dashboard_5()
    print("Generating Dashboard 6: Compliance Officer...")
    db6 = draw_dashboard_6()
    print("Generating Dashboard 7: Credit Scoring Engine...")
    db7 = draw_dashboard_7()

    # ─── Workflow Diagrams Section ───
    doc.add_heading("Section 1: Workflow Diagrams", level=1)

    workflows = [
        ("1.1 Loan Application & Disbursement Flow", wf1,
         "This flow illustrates the origination logic, showing how the system branches "
         "depending on whether the loan requires a single lump-sum payout or stage-wise tranches."),
        ("1.2 M-Pesa C2B Repayment & Auto-Clearing", wf2,
         "Maps the Safaricom Daraja webhook integration that processes incoming Paybill payments "
         "and auto-clears the loan when the outstanding balance reaches zero."),
        ("1.3 Nightly Penalties & Defaults (Cron Job)", wf3,
         "Maps the logic within scripts/daily_penalties.py which runs at midnight to catch "
         "overdue borrowers and apply the configured penalty rate."),
        ("1.4 Loan Top-Up & Settlement Flow", wf4,
         "New feature flow: existing borrowers request additional credit on active loans. "
         "The system merges balances and recalculates the amortization schedule."),
        ("1.5 Group Lending — Joint Liability Disbursement", wf5,
         "New feature flow: group lending where members share joint liability. "
         "If one member defaults, the remaining group members cover the shortfall."),
    ]

    for title, img_bytes, description in workflows:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(9.5))
        doc.add_page_break()

    # ─── Dashboard Mockups Section ───
    doc.add_heading("Section 2: Role-Based Dashboard Wireframes", level=1)
    doc.add_paragraph(
        "Each dashboard below is tailored to a specific operational role within Karibu Credit. "
        "The admin panel includes a role-preview switcher that lets privileged users view any "
        "dashboard without changing actual permissions."
    )

    dashboards = [
        ("2.1 CEO / Managing Director", db1,
         "Executive overview with KPI metrics (Total Portfolio, Revenue, Profit Margin, Active Customers), "
         "disbursement vs collection bar chart, and portfolio product mix doughnut."),
        ("2.2 CFO / Finance", db2,
         "Financial dashboard showing revenue streams (Interest Income, Processing Fees, Gross Profit), "
         "Daraja B2C/C2B settlement audit table, and provision reserve tracking."),
        ("2.3 Branch Manager", db3,
         "Branch-level dashboard with officer performance board (TAT, PAR 30 per officer), "
         "branch portfolio stats, and PAR aging breakdown chart."),
        ("2.4 Loan Officer", db4,
         "Operational dashboard with pending application queue, appraisal checklist panel, "
         "and approval workflow controls (Approve/Reject/Escalate)."),
        ("2.5 Collections Officer", db5,
         "Recovery-focused dashboard with Kanban-style arrears board (1-30 days, 31-60 days, 60+ days), "
         "restructuring controls, and CRB submission tracking."),
        ("2.6 Compliance Officer", db6,
         "Regulatory compliance dashboard with CBK filing tracker, AML transaction flags, "
         "and KYC verification queue with 48-hour SLA monitoring."),
        ("2.7 Credit Scoring Engine", db7,
         "Analytical dashboard with score distribution histogram, model weight tuning controls, "
         "and credit tier eligibility matrix."),
    ]

    for title, img_bytes, description in dashboards:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(9.5))
        doc.add_page_break()

    # ─── Implementation Notes ───
    doc.add_heading("Section 3: Implementation Notes", level=1)

    doc.add_heading("3.1 Clean Architecture (Dependency Rule)", level=2)
    doc.add_paragraph(
        "The Domain Layer is the inner-most layer. It must not import from Infrastructure. "
        "Application services orchestrate use cases by calling domain entities and repository ports. "
        "Infrastructure adapters (Daraja, ZamuPay, SQLAlchemy) implement these ports."
    )

    doc.add_heading("3.2 Security & CBK Compliance", level=2)
    doc.add_paragraph(
        "All endpoints enforce JWT-based RBAC. Audit trails log every sensitive action (loan approvals, "
        "borrower onboarding, CRB checks, policy changes). PAR thresholds (PAR 30 < 5%, NPL < 3%) "
        "are monitored and reported to CBK per NDTCP Regulations."
    )

    doc.add_heading("3.3 Repayment Allocation Waterfall", level=2)
    doc.add_paragraph(
        "Incoming payments are allocated in order: (1) Outstanding Penalties/Fees, "
        "(2) Scheduled Installments (interest first, then principal per line), "
        "(3) Overpayment to global principal, (4) Auto-clearing if balance ≤ 0."
    )

    # ─── Save ───
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Karibu_Credit_Technical_Design_v1.docx")

    doc.save(output_path)
    print(f"\n✅ Document saved to: {output_path}")
    print(f"   Size: {os.path.getsize(output_path) / 1024:.0f} KB")
    print(f"   Contains: 5 workflow diagrams + 7 dashboard wireframes")


if __name__ == "__main__":
    print("=" * 60)
    print("  Karibu Credit — Technical Design Document Generator")
    print("=" * 60)
    compile_document()
    print("\nDone! ✨")
